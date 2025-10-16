# import pdfplumber
# from docx import Document
# from langchain_community.document_loaders import TextLoader
# import tempfile
# import os
# import re

# def parse_file(uploaded_file):
#     ext = os.path.splitext(uploaded_file.name)[1].lower()

#     with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
#         tmp.write(uploaded_file.read())
#         tmp_path = tmp.name

#     if ext == ".pdf":
#         return parse_pdf(tmp_path)
#     elif ext == ".docx":
#         return parse_docx(tmp_path)
#     elif ext == ".txt":
#         return parse_txt(tmp_path)
#     else:
#         raise ValueError("Unsupported file type")

# def parse_pdf(file_path):
#     text = ""
#     with pdfplumber.open(file_path) as pdf:
#         for page in pdf.pages:
#             page_text = page.extract_text()
#             if page_text:
#                 text += page_text + "\n"
#     return text.strip()

# def parse_docx(file_path):
#     doc = Document(file_path)
#     return "\n".join([para.text for para in doc.paragraphs]).strip()

# def parse_txt(file_path):
#     loader = TextLoader(file_path)
#     doc = loader.load()
#     return doc.page_content.strip()

# def extract_hr_email_and_company(text):
#     match = re.search(r'[\w\.-]+@([\w\.-]+)', text)
#     if match:
#         email = match.group(0)
#         domain = match.group(1)
#         company = domain.split('.')[0].replace("-", " ").title()
#         return email, company
#     return None, None


import pdfplumber
from docx import Document
from langchain_community.document_loaders import TextLoader
import tempfile
import os
import re

def parse_file(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name

    if ext == ".pdf":
        text = parse_pdf(tmp_path)
    elif ext == ".docx":
        text = parse_docx(tmp_path)
    elif ext == ".txt":
        text = parse_txt(tmp_path)
    else:
        raise ValueError("Unsupported file type")

    return text.strip()

def parse_pdf(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text

def parse_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def parse_txt(file_path):
    loader = TextLoader(file_path)
    doc = loader.load()
    return doc.page_content

def extract_candidate_details(resume_text):
    name = None
    email = None
    phone = None
    linkedin = None

    # Email
    email_match = re.search(r'[\w\.-]+@[\w\.-]+', resume_text)
    if email_match:
        email = email_match.group(0)

    # Phone
    phone_match = re.search(r'(\+91[-\s]?)?\d{10}', resume_text)
    if phone_match:
        phone = phone_match.group(0)

    # LinkedIn
    linkedin_match = re.search(r'https?://(www\.)?linkedin\.com/in/[^\s]+', resume_text)
    if linkedin_match:
        linkedin = linkedin_match.group(0)

    # Name (first clean line)
    lines = resume_text.strip().splitlines()
    for line in lines[:10]:
        if line.strip() and not any(x in line.lower() for x in ["email", "@", "phone", "linkedin"]):
            name = line.strip()
            break

    return name or "Candidate", email or "Not Provided", phone or "Not Provided", linkedin or "Not Provided"

def extract_hr_email_and_company(text):
    match = re.search(r'[\w\.-]+@([\w\.-]+)', text)
    if match:
        email = match.group(0)
        domain = match.group(1)
        company = domain.split('.')[0].replace("-", " ").title()
        return email, company
    return None, None
